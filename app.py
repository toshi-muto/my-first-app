# 時系列アプリ_20250210_push
import os
import gradio as gr

# 数値計算に使うライブラリ
import numpy as np
import pandas as pd

# グラフを描画するライブラリ
from matplotlib import pyplot as plt
import matplotlib.dates as mdates
import seaborn as sns
sns.set()

# statsmodels
import statsmodels.api as sm
import statsmodels.tsa.api as tsa
from statsmodels.tsa.seasonal import seasonal_decompose

# sktime：グラフ描画
from sktime.utils.plotting import plot_series

# sktime：予測
from sktime.forecasting.naive import NaiveForecaster
from sktime.forecasting.trend import PolynomialTrendForecaster

# sktime：予測の評価指標
from sktime.performance_metrics.forecasting import (
    mean_absolute_scaled_error, MeanAbsoluteError,
    mean_absolute_percentage_error, mean_absolute_error
)

# sktime：予測の評価
from sktime.forecasting.model_selection import (
    temporal_train_test_split, ExpandingWindowSplitter, ForecastingGridSearchCV
)
from sktime.forecasting.model_evaluation import evaluate

# sktime：データの変換
from sktime.transformations.series.detrend import (
    Deseasonalizer, Detrender
)
from sktime.transformations.series.difference import Differencer
from sktime.transformations.series.boxcox import LogTransformer
from sklearn.preprocessing import StandardScaler
from sktime.transformations.series.adapt import TabularToSeriesAdaptor

# sktime：パイプライン
from sktime.forecasting.compose import (
    TransformedTargetForecaster, MultiplexForecaster
)
from sktime.pipeline import make_pipeline
from sktime.transformations.compose import OptionalPassthrough

# 機械学習法
import lightgbm as lgb

# 再帰的に回帰分析を実行するための関数の読み込み
from sktime.forecasting.compose import make_reduction

# 周期性やトレンドの関する特徴量を作成
from statsmodels.tsa.deterministic import DeterministicProcess

#グラフ日本語表記
import japanize_matplotlib

# グラフの日本語表記
from matplotlib import rcParams
rcParams['font.family'] = 'sans-serif'
rcParams['font.sans-serif'] = ['Hiragino Maru Gothic Pro','Yu Gothic','Meirio','Takao','IPAexGoshic','IPAPGothic'
                              ,'VL PGothic','Noto Sans CJK JP']

# EXCELに関する読み込み
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.page import PageMargins
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.drawing.image import Image

# フォームから入力された値を取得（例: Streamlit を使用）
import streamlit as st

# AIを利用する
import openai

# ワードを利用する
from docx import Document
import tempfile
from datetime import datetime


def preprocess_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt):
    file_path = folder.name
    data = pd.read_csv(file_path, encoding="cp932", low_memory=False)
    # dataのカラムをクリーニング
    data.columns = data.columns.str.strip()

    # 必要なデータ処理
    data['出荷金額'] = data['出荷金額'].str.replace('\\', '', regex=False).astype(float)
    data['単価'] = data['単価'].str.replace('\\', '', regex=False).astype(float)
    data['数量'] = data['数量'].str.replace('\\', '', regex=False).astype('int64')
    data['出荷数量'] = data['出荷数量'].str.replace('\\', '', regex=False).fillna(0).astype('int64')

    # 日付変換
    date_cols = ["納期回答日", "弊社回答日", "受注日", "出荷日", "納期"]
    for col in date_cols:
        data[col] = pd.to_datetime(data[col], errors="coerce")
        
    data.loc[data['型式'].isna(), '型式'] = "部品"
    data.loc[data['容量'].isna(), '容量'] = "-"

    # 列の文字列化と空白補完
    for col in ['型式', '容量', '極数', '接点構成', '接続方式', '操作電圧']:
        data[col] = data[col].fillna("-").astype(str)

    # フィルタ条件の動的生成
    query_conditions = [f'納期 >= "{day1}" and 納期 <= "{day2}"']
    if type:
        query_conditions.append(f'型式 == "{type}"')
    if capacity:
        query_conditions.append(f'容量 == "{capacity}"')
    if num:
        query_conditions.append(f'極数 == "{num}"')
    if contact:
        query_conditions.append(f'接点構成 == "{contact}"')
    if cmeth:
        query_conditions.append(f'接続方式 == "{cmeth}"')
    if volt:
        query_conditions.append(f'操作電圧 == "{volt}"')

    # 条件を結合してクエリを作成
    query_string = " and ".join(query_conditions)

    # データフィルタリング
    data_2k = data.query('得意先コード != 9999').copy()
    data_2k['受注金額'] = data_2k['数量'] * data_2k['単価']
    data_2k['変動費計'] = data_2k['数量'] * data_2k['材料外注計']
    data_2k['粗利益計'] = data_2k['受注金額'] - data_2k['変動費計']

    # 動的クエリに基づいてフィルタリング
    data_3 = data_2k.query(query_string)

    # 集計
    data_4 = data_3.groupby(['納期'])['数量'].sum().reset_index()
    data_4 = data_4.rename(columns={'納期': 'date', '数量': 'value'})

    # ベースデータ作成
    base_data = pd.DataFrame({'date': pd.date_range(day1, day2, freq='D')})
    data_base_1 = pd.merge(base_data, data_4, on='date', how='left')
    data_base_1['value'] = data_base_1['value'].fillna(0)
    data_base_1 = data_base_1.groupby([pd.to_datetime(data_base_1["date"]).dt.strftime("%Y-%m")]).value.sum().reset_index()
    data_base_1 = data_base_1.set_index('date')
    data_base_1['value'] = data_base_1['value'].astype('float')
    
    return data_base_1



def process_1_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt, bins):
    data_base_1 = preprocess_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt)
    # kishuの作成
    kishu = str(type) + str(capacity) + str(num) + str(contact) + str(cmeth) + str(volt)

    # インデックスをリセットして日付を表示データに含める
    data_display_1 = data_base_1.copy().reset_index()

    #kishu = kishu_input  # 入力された機種別の値
    file_path_1 = 'value_{}.xlsx'.format(kishu)  # format関数を使用してファイル名を生成
    data_display_1.to_excel(file_path_1, sheet_name='value', index=False) # データをExcelに保存

    # Excelファイルを開いて横向き設定を適用
    workbook = load_workbook(file_path_1)
    sheet = workbook['value']

    # シートを横向きに設定
    sheet.page_setup.orientation = 'landscape'  # 横向き設定
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A3  # 用紙サイズをA3に設定
    sheet.page_margins = PageMargins(
        left=0.5, right=0.5, top=0.5, bottom=0.5, header=0.3, footer=0.3
    )  # ページ余白を設定

    # 保存
    workbook.save(file_path_1)

    # グラフ作成_1
    plt.figure()
    data_base_1['value'].plot()
    plt.title('Monthly Value')
    plt.xlabel('Date')
    plt.ylabel('Value')
    plt.xticks(rotation=0)
    plt.tight_layout()
    plot_file_11 = "plot_11.png"
    plt.savefig(plot_file_11)
    plt.close()

    # グラフ作成_2
    plt.figure()
    data_base_1[[ 'value']].describe()
    data_base_1['value'].hist(bins=bins)
    plot_file_12 = "plot_12.png"
    plt.savefig(plot_file_12)
    plt.close()


    # Excelファイルを開く
    wb = load_workbook(file_path_1)
    workbook = openpyxl.load_workbook(file_path_1)
    sheet = workbook.active


    # 画像を挿入
    img_1 = Image("plot_11.png")
    img_1.width, img_1.height = img_1.width * 1.1, img_1.height * 1.0
    sheet.add_image(img_1, "D1")
    img_2 = Image("plot_12.png")
    img_2.width, img_2.height = img_2.width * 1.0, img_2.height * 1.0
    sheet.add_image(img_2, "D30")

    # Excelファイルを保存
    workbook.save(file_path_1)

    return plot_file_11, plot_file_12, data_display_1, file_path_1


def process_2_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt, bins):
    data_base_1 = preprocess_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt)

    # kishuの作成
    kishu = str(type) + str(capacity) + str(num) + str(contact) + str(cmeth) + str(volt)

    # データをコピー
    data_base_1_diff = data_base_1.copy()
    # 差分系列
    data_base_1_diff['diff1'] = data_base_1_diff['value'].diff(1)
    # 季節差分
    data_base_1_diff['diff12'] = data_base_1['value'].diff(12)
    # 季節差分系列に対して、さらに差分をとる
    #data_base_1_diff['diff12-1'] = data_base_1_diff['diff12'].diff(1)

    # インデックスをリセットして日付を表示データに含める
    data_display_2 = data_base_1_diff.copy().reset_index()
    #kishu = kishu_input  # 入力された機種別の値
    file_path_2 = 'diff_{}.xlsx'.format(kishu)  # format関数を使用してファイル名を生成
    data_display_2.to_excel(file_path_2, sheet_name='diff', index=False) # データをExcelに保存

    # Excelファイルを開いて横向き設定を適用
    workbook = load_workbook(file_path_2)
    sheet = workbook['diff']

    # シートを横向きに設定
    sheet.page_setup.orientation = 'landscape'  # 横向き設定
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A3  # 用紙サイズをA3に設定
    sheet.page_margins = PageMargins(
        left=0.5, right=0.5, top=0.5, bottom=0.5, header=0.3, footer=0.3
    )  # ページ余白を設定

    # 保存
    workbook.save(file_path_2)


    # グラフ作成_1
    plt.figure()
    data_base_1_diff[['diff1', 'diff12']].plot()
    plt.title('Monthly Value')
    plt.xlabel('Date')
    plt.ylabel('Value')
    plt.xticks(rotation=0)
    plt.tight_layout()
    plot_file_21 = "plot_21.png"
    plt.savefig(plot_file_21)
    plt.close()
    '''
    # グラフ作成_2（ケース１：hist）
    plt.figure()
    data_base_1_diff[['diff1']].describe()
    data_base_1_diff['diff1'].hist(bins=bins)
    plot_file_22 = "plot_22.png"
    plt.savefig(plot_file_22)
    plt.close()

    # グラフ作成_2（ケース２：diff12 相関）
    plt.figure()
    sns.scatterplot(x=data_base_1_diff['diff12'].shift(12),
                    y=data_base_1_diff['diff12'])
    plot_file_22 = "plot_22.png"
    plt.savefig(plot_file_22)
    plt.close()
    '''
    # グラフ作成_3（ケース３：seasonal_decompose）
    data_base_10 = data_base_1.copy()
    if not isinstance(data_base_10.index, pd.DatetimeIndex):
        data_base_10.index = pd.to_datetime(data_base_10.index)
    plt.figure()
    plt.rcParams['figure.figsize'] = [8, 6]  # グラフサイズ
    plt.rcParams['font.size'] = 10  # フォントサイズ
    result = seasonal_decompose(data_base_10.value, model='additive', period=12).plot()
    plot_file_22 = "plot_22.png"
    plt.savefig(plot_file_22)
    plt.close()


    # Excelファイルを開く
    wb = load_workbook(file_path_2)
    sheet = workbook.active

    # 画像を挿入
    img_1 = Image("plot_21.png")
    img_1.width, img_1.height = img_1.width * 1.0, img_1.height * 1.0
    sheet.add_image(img_1, "G1")
    img_2 = Image("plot_22.png")
    img_2.width, img_2.height = img_2.width * 1.0, img_2.height * 1.0
    sheet.add_image(img_2, "G37")

    # Excelファイルを保存
    workbook.save(file_path_2)

    return plot_file_21, plot_file_22, data_display_2, file_path_2



def process_3_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt, x_cson, test_size):
    data_base_1 = preprocess_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt)

    # kishuの作成
    kishu = str(type) + str(capacity) + str(num) + str(contact) + str(cmeth) + str(volt)

    data_base_5 = data_base_1.copy()

    # インデックスをリセットして日付を表示データに含める
    # または PeriodIndex に変換
    data_base_5.index = pd.period_range(start=data_base_5.index[0], periods=len(data_base_5), freq='M')

    # 訓練データとテストデータに分割
    train_data, test_data = temporal_train_test_split(data_base_5, test_size=test_size)

    # 予測期間
    fh_data = np.arange(1, len(test_data) + 1)

    # 加重平均の重み
    weight = np.concatenate([np.array([1/8]), np.tile(1/4, 3), np.array([1/8])])

    # データをコピー
    data_base_1_ma = data_base_5.copy()

    # 再コピー（警告回避のため）
    data_base_1_ma = data_base_1_ma.copy()

    # 5時点の移動平均
    data_base_1_ma['ma5'] = data_base_1_ma['value'].rolling(
        window=5).mean()

    # 5時点の移動平均において、当該時点を中心にする
    data_base_1_ma['ma5_center'] = data_base_1_ma['value'].rolling(
        window=5, center=True).mean()

    # 4時点移動平均
    data_base_1_ma['ma4_center'] = data_base_1_ma['value'].rolling(
        window=4, center=True).mean()

    # 加重平均をまとめて計算
    true_ma4 = data_base_1_ma['value'].rolling(
        window=5, center=True).apply(np.average, kwargs={'weights': weight})

    # 単純な4月移動平均に対して、さらに移動平均をとる
    data_base_1_ma['true_ma4'] = \
        data_base_1_ma['ma4_center'].rolling(window=2).mean().shift(-1)

    # 12時点移動平均
    ma_12 = data_base_5['value'].rolling(window=12, center=True).mean()

    # 単純な12時点移動平均に対して、さらに移動平均をとる
    # これが中心化移動平均となる
    data_base_1_ma['trend'] = ma_12.rolling(window=2).mean().shift(-1)
    data_base_1_ma['trend'] = data_base_1_ma['trend'].fillna(data_base_1_ma['trend'].median())
    data_base_1_ma['true_ma4'] = data_base_1_ma['true_ma4'].fillna(data_base_1_ma['trend'].median())
    data_base_1_ma['ma5_center'] = data_base_1_ma['ma5_center'].fillna(data_base_1_ma['trend'].median())
    data_base_1_ma['ma4_center'] = data_base_1_ma['ma4_center'].fillna(data_base_1_ma['trend'].median())

    # 対数変換を適用
    data_base_1_ma['value_log'] = np.log(data_base_1_ma['value'] + 1e-9)
    data_base_1_ma['true_ma4_log'] = np.log(data_base_1_ma['true_ma4'] + 1e-9)
    data_base_1_ma['trend_log'] = np.log(data_base_1_ma['trend'] + 1e-9)
    data_base_1_ma['ma5_center_log'] = np.log(data_base_1_ma['ma5_center'] + 1e-9)
    data_base_1_ma['ma4_center_log'] = np.log(data_base_1_ma['ma4_center'] + 1e-9)
    data_base_1_mat = data_base_1_ma[['value_log','true_ma4_log', 'trend_log']]
    data_base_1_mat.index = pd.period_range(start=data_base_1_mat.index[0], periods=len(data_base_1_mat), freq='M')
    #data_base_1_mat.index = pd.period_range(start=data_base_1_mat.index[0], periods=len(data_base_1_mat), freq='M').to_timestamp()

    
    X_cs = x_cson

    X_train = pd.DataFrame()
    X_test = pd.DataFrame()
    Xa_train = pd.DataFrame()
    Xa_test = pd.DataFrame()
    Xc_train = pd.DataFrame()
    Xc_test = pd.DataFrame()

    # option_1 特微量は移動平均のみ
    if X_cs == "option1":
        X_train = data_base_1_mat.loc[train_data.index]
        X_test = data_base_1_mat.loc[test_data.index]
    
    # option_2 特微量は三角関数のみ こちらを復活させる場合はoptionに2を追加する
    elif X_cs == "option2":
        Xa_train = data_base_1_mat.loc[train_data.index, 'value_log']
        Xa_test = data_base_1_mat.loc[test_data.index, 'value_log']
        
        dp = DeterministicProcess(
            train_data.index, constant=False, order=1, period=24, fourier=6
        )
        # 訓練データ
        Xc_train = dp.in_sample()
        # 訓練データの列名の変更
        Xc_train.columns = ["trend_1"] + [f"seasonal_{x}" for x in range(6 * 2)]
        # テストデータ
        Xc_test = dp.out_of_sample(len(test_data))
        # テストデータの列名の変更
        Xc_test.columns = ["trend_1"] + [f"seasonal_{x}" for x in range(6 * 2)]
        # 訓練データを結合
        X_train = pd.concat([Xa_train, Xc_train], axis=1)
        # テストデータを結合
        X_test = pd.concat([Xa_test, Xc_test], axis=1)
        
    # option_3 特微量は移動平均と三角関数の両方
    elif X_cs == "option3":
        Xa_train = data_base_1_mat.loc[train_data.index]
        Xa_test = data_base_1_mat.loc[test_data.index]

        dp = DeterministicProcess(
            train_data.index, constant=False, order=1, period=24, fourier=6
        )
        # 訓練データ
        Xc_train = dp.in_sample()
        # 訓練データの列名の変更
        Xc_train.columns = ["trend_1"] + [f"seasonal_{x}" for x in range(6 * 2)]
        # テストデータ
        Xc_test = dp.out_of_sample(len(test_data))
        # テストデータの列名の変更
        Xc_test.columns = ["trend_1"] + [f"seasonal_{x}" for x in range(6 * 2)]
        # 訓練データを結合
        X_train = pd.concat([Xa_train, Xc_train], axis=1)
        # テストデータを結合
        X_test = pd.concat([Xa_test, Xc_test], axis=1)       
        


    # LightGBMのハイパーパラメータを設定
    params_3 = {
        'objective': 'regression',    # 回帰を目的とする
        'seed': 1,                    # 乱数の種
        'num_leaves': 60,             # 葉の数の最大値
        'learning_rate': 0.07,        # 学習率
        'n_estimators': 100,          # ブースティングの回数
        'min_data_in_leaf': 4,        # 1つの葉における最小データ
        'verbose': -1                 # ワーニングなどの非表示
    }

    # モデル化
    gbm_sk_data = lgb.LGBMRegressor(**params_3)

    # 再帰的にLightGBMを実行
    gbm_forecaster_data = make_reduction(
        gbm_sk_data, window_length=12, pooling='global', strategy="recursive")

    # モデルの当てはめ
    gbm_forecaster_data.fit(train_data)

    detrend = Detrender(
    forecaster=PolynomialTrendForecaster(degree=1),
    model='multiplicative')

    # 前処理からモデル化までを1つのパイプラインにまとめる
    pipe_gbm = TransformedTargetForecaster(
        [
            detrend,
            ('forecast', make_reduction(
                gbm_sk_data, window_length=12, pooling='global',
                strategy="recursive")),
        ]
    )

    # データへの当てはめ
    pipe_gbm.fit(train_data)
    # 前処理
    transed = detrend.fit_transform(train_data)
    # モデル化
    gbm_reg = lgb.LGBMRegressor(**params_3)
    # モデルの当てはめ
    gbm_reg.fit(X_train, transed)
    # 予測の実施
    pred_gbm_reg = pd.DataFrame(
        {'value':gbm_reg.predict(X_test)}, index=test_data.index)
    # 変換をもとに戻す
    pred_gbm_reg = detrend.inverse_transform(pred_gbm_reg)

    # LightGBMモデル設定
    gbm_sk_data = lgb.LGBMRegressor(**params_3)
    # LightGBMに渡すデータをコピー
    X_train = X_train.copy()
    X_test = X_test.copy()
    train_data = train_data.copy()
    test_data = test_data.copy()

    # LightGBMモデルの学習
    gbm_sk_data.fit(X_train, train_data)
    # 予測の実施
    ma4_pred_ma = pd.DataFrame({'value': gbm_sk_data.predict(X_test)}, index=test_data.index)
    # 評価指標の計算
    mae = mean_absolute_error(test_data, ma4_pred_ma)
    mase = mean_absolute_scaled_error(test_data, ma4_pred_ma, y_train=train_data)
    
    # 特微量を X_cs に応じて設定
    if X_cs == "option1":
        # value_log を先頭に配置
        long_term_exog_2 = X_test[['value_log', 'true_ma4_log', 'trend_log']].astype(float).values
        X_train_selected = X_train[['value_log', 'true_ma4_log', 'trend_log']]

    elif X_cs == "option2":
        # 三角関数のみ
        long_term_exog_2 = X_test[['value_log', 'trend_1'] + [f'seasonal_{i}' for i in range(12)]].astype(float).values
        X_train_selected = X_train[['value_log', 'trend_1'] + [f'seasonal_{i}' for i in range(12)]]

    elif X_cs == "option3":
        # value_log を先頭にして、他の列を追加
        long_term_exog_2 = X_test[['value_log', 'true_ma4_log', 'trend_log', 'trend_1'] + [f'seasonal_{i}' for i in range(12)]].astype(float).values
        X_train_selected = X_train[['value_log', 'true_ma4_log', 'trend_log', 'trend_1'] + [f'seasonal_{i}' for i in range(12)]]
    
    future_index = pd.date_range(
        start=test_data.index[-1].to_timestamp() + pd.offsets.MonthEnd(1),
        periods=test_size,
        freq='MS'
    )
    
    # モデルの再学習
    gbm_reg = lgb.LGBMRegressor(**params_3)
    gbm_reg.fit(X_train_selected.iloc[:, 1:], X_train_selected.iloc[:, 0])  # 説明変数: 列1以降, 目的変数: 列0

    # 長期予測の値を格納するリスト
    long_term_forecast_2 = []

    # 長期予測を逐次生成する
    for i in range(test_size):
        # 説明変数は列1以降（列0は目的変数）
        X_current = long_term_exog_2[i, 1:].reshape(1, -1)

        # モデルで予測
        forecast_value_2 = gbm_reg.predict(X_current)[0]
        long_term_forecast_2.append(forecast_value_2)

        # 次の特徴量を更新
        if i < test_size - 1:
            long_term_exog_2[i + 1, 0] = forecast_value_2  # 列0（value_log）を更新

    # 長期予測結果のデータフレームを作成（逆変換を適用）
    forecast_df = pd.DataFrame({
        'Forecast': np.exp(long_term_forecast_2)  # 自然対数の逆変換
    }, index=future_index)

    # 必要に応じて、表示や出力時に PeriodIndex に変換　ケース２
    forecast_df.index = forecast_df.index.to_period('M')
    
    # データの準備と結合処理
    p1 = ma4_pred_ma.copy()
    p1 = p1.rename(columns={'value': 'pred_2'})
    # テストデータと予測データを結合
    g = pd.concat([test_data, p1], axis=1)
    # 予測データフレームを作成
    #forecast_df = pd.DataFrame({
    #    'Forecast': long_term_forecast_2  # 予測値を列に設定
    #}, index=forecast_df.index)  # `future_index`をインデックスとして設定　forecast_df.indexか？

    # すべてのデータを結合
    g_2 = pd.concat([g, forecast_df], axis=1)
    # 小数点第二位まで丸める
    data_display_3 = g_2.round(2).reset_index()

    #kishu = kishu_input  # 入力された機種別の値
    file_path_3 = 'forecast_{}.xlsx'.format(kishu)  # format関数を使用してファイル名を生成

    #data_display_3.to_csv(file_path_3, index=False) #csvにする場合
    data_display_3.to_excel(file_path_3, sheet_name='forecast', index=False) # データをExcelに保存

    # Excelファイルを開いて横向き設定を適用
    workbook = load_workbook(file_path_3)
    sheet = workbook['forecast']
    # シートを横向きに設定
    sheet.page_setup.orientation = 'landscape'  # 横向き設定
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A3  # 用紙サイズをA3に設定
    sheet.page_margins = PageMargins(
        left=0.5, right=0.5, top=0.5, bottom=0.5, header=0.3, footer=0.3
    )  # ページ余白を設定

    # 保存
    workbook.save(file_path_3)

    # グラフを保存するファイル名を定義
    # 日本語フォントを指定
    rcParams['font.family'] = 'IPAexGothic'
    # グラフ作成
    plt.figure(figsize=(8, 4))

    # 予測結果の可視化
    plot_file_31 = "plot_31.png"  # ここでファイル名を定義
    fig, ax = plot_series(
        train_data, test_data, ma4_pred_ma,
        labels=['train', 'test', 'test_pred'],
        markers=np.tile('', 3)
    )

    # 長期予測をプロット
    ax.plot(
        forecast_df.index.to_timestamp(), forecast_df['Forecast'],  # 正しいスケールの値をプロット
        label='長期予測 (24ステップ)', linestyle='-', color='red'
    )

    # グラフの装飾
    ax.legend(title='予測結果', fontsize=10, title_fontsize=12)
    # グラフの外にテキストボックスとして表示
    plt.text(0.05, 1.05, f"MASE: {mase:.2f}", transform=ax.transAxes, fontsize=9)
    plt.text(0.05, 1.09, f"MAE: {mae:.2f}", transform=ax.transAxes, fontsize=9)

    # 予測時系列図_1
    plt.tight_layout()
    plt.xticks(rotation=0)
    plt.savefig(plot_file_31)  # ファイルを保存
    plt.close(fig)

    # 予測時系列図_2
    
    #future_index = pd.date_range(start=test_data.index[-1].to_timestamp(), periods=test_size, freq='MS')
    plt.figure(figsize=(8, 4))
    plt.plot(forecast_df.index.to_timestamp(), forecast_df['Forecast'], label='長期予測 (24ステップ)', linestyle='-', color='red')
    #plt.plot(forecast_df.index.to_timestamp(), long_term_forecast_2, label='長期予測', linestyle='--', color='red')
    #plt.plot(forecast_df.index, long_term_forecast_2, label='長期予測', linestyle='--', color='red')
    plt.legend()
    plt.title('長期予測の結果', fontsize=14)
    plt.xlabel('年月', fontsize=12)
    plt.ylabel('値', fontsize=12)
    plt.xticks(rotation=0)
    plot_file_32 = "plot_32.png"  # ここでファイル名を定義
    plt.savefig(plot_file_32)
    plt.close()
    
    # 予測図のEXCELへの挿入
    # Excelファイルを開く
    wb = load_workbook(file_path_3)
    workbook = openpyxl.load_workbook(file_path_3)
    sheet = workbook.active

    # 画像を挿入
    img_1 = Image("plot_31.png")
    img_1.width, img_1.height = img_1.width * 0.7, img_1.height * 1.2
    sheet.add_image(img_1, "F1")
    img_2 = Image("plot_32.png")
    img_2.width, img_2.height = img_2.width * 0.8, img_2.height * 0.8
    sheet.add_image(img_2, "F29")

    # Excelファイルを保存
    workbook.save(file_path_3)

    # 予測結果のデータフレームとグラフのパスを返す
    return plot_file_31, plot_file_32, data_display_3, file_path_3


# GPTによる応答生成（データをプロンプトに含める）
def generate_gpt_response(folder, day1, day2,
                          type, capacity, num, contact, cmeth, volt, use_chatgpt, token, gpt_model, prompt):

    if not use_chatgpt:
        # ChatGPTを使用しない場合の処理
        return "ChatGPTを使用しない設定のため、処理をスキップしました。", None
    

    # データを処理
    data_base_11 = preprocess_data(folder, day1, day2, type, capacity, num, contact, cmeth, volt)
    # kishuの作成
    kishu = str(type) + str(capacity) + str(num) + str(contact) + str(cmeth) + str(volt)

    data_display_1c = data_base_11.copy()
    data_display_1c = data_display_1c.reset_index()

    openai.api_key = token

    # プロンプトテンプレートを定義
    formatted_prompt = f"{prompt}\n\n{data_display_1c.to_string(index=False)}"

    try:
        # GPTに質問を投げる
        response = openai.chat.completions.create(
            model=gpt_model,
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": formatted_prompt}
            ]
        )

        gpt_response = response.choices[0].message.content

        # 今日の日付を取得し、フォーマット
        today = datetime.now().strftime("%Y-%m-%d")

        # 一時ファイルを作成（ローカルで行う場合）
        tmp_dir = tempfile.gettempdir()
        file_path_cat_1 = f"{tmp_dir}\\ChatGPT_{kishu}.docx"
        
        # 一時ファイルを作成（外部で行う場合）
        #tmp_dir = "/tmp"
        #file_path_cat_1 = os.path.join(tmp_dir, f"ChatGPT_{kishu}.docx")
            
        print(f"ファイルが作成されました: {file_path_cat_1}")

        # ファイルに応答内容を保存
        save_report_1(gpt_response, file_path_cat_1)

        return file_path_cat_1  # Gradioのfile_download_2に渡すパスとして返す

    except Exception as e:
        print(f"エラーが発生しました: {e}")
        return None

def save_report_1(content, file_path_cat_1):
    try:
        doc = Document()
        doc.add_heading("ChatGPTの見解", level=1)
        doc.add_paragraph(content)
        doc.save(file_path_cat_1)
    except Exception as e:
        print(f"ファイル保存中にエラーが発生しました: {e}")


# Gradioインターフェースの構築
def create_interface():
    with gr.Blocks() as app:
        with gr.Row():
            # 左側の入力列
            with gr.Column(scale=1):
                gr.Markdown("Var.4.2_20250209")  # ヘッダーを追加
                folder_input = gr.File(label="CSVファイルをアップロード", file_types=[".csv"])

                # 開始日と終了日を横並びに配置
                with gr.Row():
                    day1_input = gr.Textbox(label="開始日 (day1)", value="2016-3-1")
                    day2_input = gr.Textbox(label="終了日 (day2)", value="2024-12-31")

                    # test_size_inputを正しく定義
                    test_size_input = gr.Number(
                        label="testdata & forecast（月数）基本24",
                        value=24,  # デフォルト値
                        minimum=1,
                        maximum=48
                    )

                # ここでupdate_dropdown関数を定義
                #kishu_input = gr.Textbox(label="機種別 (kishu)", value="TGMEs10")
                # パラメータ入力エリア
                with gr.Row():
                    with gr.Column():
                        # 1行目: 型式・容量
                        with gr.Row():
                            type_toggle = gr.Checkbox(label="型式を指定する", value=True)
                            type_input = gr.Textbox(label="型式 (type)", value="TGMEs", interactive=True)

                            capacity_toggle = gr.Checkbox(label="容量を指定する", value=True)
                            capacity_input = gr.Number(label="容量 (capacity)", value=10, interactive=True)

                        # 2行目: 極数・接点構成
                        with gr.Row():
                            num_toggle = gr.Checkbox(label="極数を指定する", value=False)
                            num_input = gr.Number(label="極数 (num)", value=None, interactive=False)

                            contact_toggle = gr.Checkbox(label="接点構成を指定する", value=False)
                            contact_input = gr.Textbox(label="接点構成 (contact)", value="", interactive=False)

                            cmeth_toggle = gr.Checkbox(label="接続方式を指定する", value=False)
                            cmeth_input = gr.Textbox(label="接続方式 (cmeth)", value="", interactive=False)

                            volt_toggle = gr.Checkbox(label="操作電圧を指定する", value=False)
                            volt_input = gr.Textbox(label="操作電圧 (volt)", value="", interactive=False)

                # トグルボタンで入力の有効/無効を切り替えるロジック
                def toggle_inputs(toggle, current_value):
                    return gr.update(interactive=toggle, value=current_value if toggle else None)

                # トグルの変更イベントを登録
                type_toggle.change(toggle_inputs, inputs=[type_toggle, type_input], outputs=type_input)
                capacity_toggle.change(toggle_inputs, inputs=[capacity_toggle, capacity_input], outputs=capacity_input)
                num_toggle.change(toggle_inputs, inputs=[num_toggle, num_input], outputs=num_input)
                contact_toggle.change(toggle_inputs, inputs=[contact_toggle, contact_input], outputs=contact_input)
                cmeth_toggle.change(toggle_inputs, inputs=[cmeth_toggle, cmeth_input], outputs=cmeth_input)
                volt_toggle.change(toggle_inputs, inputs=[volt_toggle, volt_input], outputs=volt_input)

                x_cson_input = gr.Dropdown(
                    label="特徴量選択",
                    choices=["option1", "option2", "option3"],
                    value="option1"
                )

                # ボタンの配置
                button_1 = gr.Button("データを処理")
                button_2 = gr.Button("データをdiff処理")
                button_3 = gr.Button("時系列予測")

                # ChatGPT関連の設定
                # チェックボックスの状態に応じて入力フィールドを切り替え
                def toggle_chatgpt_inputs(use_chatgpt):
                    if use_chatgpt:
                        # 使用可能にする
                        return (
                            gr.update(interactive=True),  # gpt_model_input
                            gr.update(interactive=True),  # prompt_input
                            gr.update(interactive=True),  # chatgpt_button
                            gr.update(interactive=True),  # token_input
                        )
                    else:
                        # 使用不可にする
                        return (
                            gr.update(interactive=False),  # gpt_model_input
                            gr.update(interactive=False),  # prompt_input
                            gr.update(interactive=False),  # chatgpt_button
                            gr.update(interactive=False),  # token_input
                        )

                with gr.Row():
                    use_chatgpt_toggle = gr.Checkbox(label="ChatGPTを使用する", value=False)
                    gpt_model_input = gr.Dropdown(
                        label="ChatGPTモデル選択",
                        choices=["gpt-3.5-turbo", "gpt-4o"],
                        value="gpt-4o",
                        interactive=False  # 初期状態で無効化
                    )

                    prompt_input = gr.Textbox(
                        label="ChatGPTへの質問",
                        value="以下の時系列データを元にこれまでの傾向\n条件1:年単位\n条件2:月単位\n条件3:異常値\n条件4:今後の予測は上昇傾向か下降傾向か",
                        lines=3,
                        placeholder="質問を入力してください",
                        interactive=False,  # 初期状態で無効化
                        scale=2
                    )
                    chatgpt_button = gr.Button("ChatGPT実行", interactive=False)
                    token_input = gr.Textbox(
                        label="ChatGPT APIトークン",
                        value="",
                        placeholder="APIキーを入力してください",
                        type="password",  # パスワード入力
                        interactive=False,  # 初期状態で無効化
                        #scale=0.5,  # ボックス幅を小さく
                        #lines=0.1     # 高さの調整（行数）
                    )

                    # チェックボックスの変更を監視して処理を実行
                    use_chatgpt_toggle.change(
                        toggle_chatgpt_inputs,
                        inputs=[use_chatgpt_toggle],
                        outputs=[gpt_model_input, prompt_input, chatgpt_button, token_input],
                    )


            # 右側の出力エリア
            with gr.Column(scale=2):
                plot_output_1 = gr.Image(label="時系列グラフ")  # 上部に時系列グラフを配置

                # データフレームと別のグラフを横並びに配置
                with gr.Row():
                    data_output = gr.Dataframe(label="データ")  # 左側にデータ
                    plot_output_2 = gr.Image(label="別のグラフ")  # 右側に別のグラフ

                # データダウンロードを別のグラフの下に配置
                bins_input = gr.Number(label="ヒストグラムのbinsを指定", value=30, minimum=5)
                file_download = gr.File(label="データダウンロード")
                file_download_2 = gr.File(label="ChatGPT報告書ダウンロード")


        # ボタンのクリック処理
        button_1.click(
            fn=process_1_data,
            inputs=[folder_input, day1_input, day2_input,
                    type_input, capacity_input, num_input, contact_input, cmeth_input, volt_input, bins_input],
            outputs=[plot_output_1, plot_output_2, data_output, file_download],
        )
        
        button_2.click(
            fn=process_2_data,
            inputs=[folder_input, day1_input, day2_input,
                    type_input, capacity_input, num_input, contact_input, cmeth_input, volt_input, bins_input],
            outputs=[plot_output_1, plot_output_2, data_output, file_download],
        )

        button_3.click(
            fn=process_3_data,
            inputs=[folder_input, day1_input, day2_input,
                    type_input, capacity_input, num_input, contact_input, cmeth_input, volt_input,
                    x_cson_input, test_size_input],
            outputs=[plot_output_1, plot_output_2, data_output, file_download],
        )

        chatgpt_button.click(
            fn=generate_gpt_response,
            inputs=[folder_input, day1_input, day2_input,
                    type_input, capacity_input, num_input, contact_input, cmeth_input, volt_input,
                    use_chatgpt_toggle, token_input, gpt_model_input, prompt_input],
            outputs=[file_download_2]
        )

    return app

app = create_interface()
app.launch()